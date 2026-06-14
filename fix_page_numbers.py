# -*- coding: utf-8 -*-
import sys, io, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

FOLDER = os.path.dirname(os.path.abspath(__file__))
FILE_NAME = "66PM_0182166_LaiXuanChinh_DATN66PM-2026.docx"
INPUT_PATH = os.path.join(FOLDER, "Báo cáo", FILE_NAME)
OUTPUT_PATH = os.path.join(FOLDER, "Báo cáo", FILE_NAME.replace(".docx", "_numbered.docx"))

doc = Document(INPUT_PATH)
print(f"Loaded. Sections={len(doc.sections)}, Paragraphs={len(doc.paragraphs)}")

# ── 1. Tìm đoạn "MỞ ĐẦU" thật (bỏ qua dòng mục lục có tab+số trang) ──────────
mo_dau_idx = None
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    # Dòng mục lục dạng "MỞ ĐẦU\t1" — bỏ qua
    if ("MỞ ĐẦU" in t or "MO DAU" in t.upper()) and '\t' not in t:
        print(f"  Found at [{i}] style={para.style.name!r}: {t[:60]!r}")
        mo_dau_idx = i
        break

if mo_dau_idx is None:
    print("KHÔNG TÌM THẤY MỞ ĐẦU — in tất cả đoạn chứa 'MỞ':")
    for i, p in enumerate(doc.paragraphs):
        if 'MỞ' in p.text or 'MO' in p.text.upper():
            print(f"  [{i}] {p.style.name}: {p.text.strip()[:60]!r}")
    sys.exit(1)

# ── 2. Helper: tạo footer có số trang căn giữa ────────────────────────────────
def make_footer_with_page_number(section):
    """Thêm số trang căn giữa vào footer của section."""
    footer = section.footer
    footer.is_linked_to_previous = False

    # Xóa nội dung cũ
    for para in footer.paragraphs:
        p_elem = para._p
        p_elem.getparent().remove(p_elem)

    # Tạo đoạn mới
    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    r_elem = run._r

    # <w:fldChar w:fldCharType="begin"/>
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r_elem.append(fldChar_begin)

    # <w:instrText> PAGE </w:instrText>
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    r_elem.append(instrText)

    # <w:fldChar w:fldCharType="end"/>
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r_elem.append(fldChar_end)


def clear_footer(section):
    """Xóa footer (không hiển thị số trang)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ''


def set_section_page_number_start(section, start=1):
    """Đặt số trang bắt đầu cho section."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start))
    pgNumType.set(qn('w:fmt'), 'decimal')


# ── 3. Chèn section break trước đoạn "MỞ ĐẦU" ────────────────────────────────
mo_dau_para = doc.paragraphs[mo_dau_idx]
mo_dau_p_elem = mo_dau_para._p

# Lấy sectPr của section hiện tại (body-level sectPr hoặc tạo mới)
body = doc.element.body
last_sectPr = body.find(qn('w:sectPr'))  # sectPr cuối cùng trong body

# Tạo sectPr mới để đặt TRƯỚC đoạn MỞ ĐẦU (kết thúc section 1)
new_sectPr = OxmlElement('w:sectPr')

# Đặt kiểu ngắt section
pgSz_elem = last_sectPr.find(qn('w:pgSz'))
if pgSz_elem is not None:
    new_sectPr.append(copy.deepcopy(pgSz_elem))

pgMar_elem = last_sectPr.find(qn('w:pgMar'))
if pgMar_elem is not None:
    new_sectPr.append(copy.deepcopy(pgMar_elem))

# Chèn pPr/sectPr vào đoạn ngay TRƯỚC MỞ ĐẦU để tạo section break
prev_para = doc.paragraphs[mo_dau_idx - 1] if mo_dau_idx > 0 else None
if prev_para:
    pPr = prev_para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        prev_para._p.insert(0, pPr)
    # Xóa sectPr cũ trong pPr nếu có
    old = pPr.find(qn('w:sectPr'))
    if old is not None:
        pPr.remove(old)
    pPr.append(new_sectPr)
    print(f"  Section break inserted before paragraph [{mo_dau_idx}]")

# ── 4. Cấu hình sections sau khi chèn break ────────────────────────────────────
# Reload sections
doc_sections = doc.sections
print(f"  Sections after insert: {len(doc_sections)}")

# Section 0 (trước MỞ ĐẦU): không có số trang
clear_footer(doc_sections[0])

# Section 1 (từ MỞ ĐẦU trở đi): số trang bắt đầu từ 1, căn giữa
if len(doc_sections) > 1:
    make_footer_with_page_number(doc_sections[1])
    set_section_page_number_start(doc_sections[1], start=1)
    print("  Footer with page number added to section 1 (from MỞ ĐẦU)")
else:
    # Chỉ có 1 section — thêm thẳng vào
    make_footer_with_page_number(doc_sections[0])
    set_section_page_number_start(doc_sections[0], start=1)
    print("  Warning: only 1 section found, added to section 0")

# ── 5. Lưu file ──────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"\nDone! Saved to:\n  {OUTPUT_PATH}")
print("\nLƯU Ý: Mở file trong Word rồi nhấn Ctrl+A → F9 để cập nhật mục lục.")
